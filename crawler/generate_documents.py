#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate test documents for crawler testing."""

import os
import random
from pathlib import Path
from datetime import datetime

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
import zipfile

try:
    import py7zr
    SEVENZIP_AVAILABLE = True
except ImportError:
    SEVENZIP_AVAILABLE = False

try:
    import rarfile
    RARFILE_AVAILABLE = True
except ImportError:
    RARFILE_AVAILABLE = False

WORDS = [
    "document", "file", "data", "system", "program", "computer", "server", "database",
    "network", "internet", "application", "user", "security", "access", "report",
    "analysis", "table", "list", "search", "filter", "export", "import", "backup",
    "module", "library", "technology", "development", "testing", "project", "task",
    "meeting", "presentation", "training", "chapter", "section", "text", "content"
]

TOPICS = [
    "Annual Report", "Financial Statement", "Project Plan", "Technical Specification",
    "User Manual", "System Documentation", "Research Analysis", "Marketing Strategy",
    "Employee Handbook", "Quality Control", "Sales Report", "Inventory List"
]

COMPANIES = [
    "TechCorp Solutions", "DataFlow Systems", "CloudNet Services", "InfoTech Industries",
    "Digital Innovations", "Global Solutions Inc"
]


def random_sentence(min_words=5, max_words=15):
    """Generate random sentence."""
    words = [random.choice(WORDS) for _ in range(random.randint(min_words, max_words))]
    return " ".join(words).capitalize() + "."


def random_paragraph(min_sentences=3, max_sentences=7):
    """Generate random paragraph."""
    return " ".join(random_sentence() for _ in range(random.randint(min_sentences, max_sentences)))


def random_table(rows=5, cols=4):
    """Generate random table data."""
    data = [["Column " + str(i+1) for i in range(cols)]]
    for _ in range(rows):
        row = [f"Item {random.randint(1000, 9999)}"]
        for j in range(1, cols-1):
            row.append(random_sentence(2, 4))
        row.append(f"{random.randint(100, 9999)}.{random.randint(0, 99):02d}")
        data.append(row)
    return data


def create_docx(path, title=None):
    """Create DOCX document."""
    if title is None:
        title = random.choice(TOPICS)
    
    doc = Document()
    doc.add_heading(title, 0).alignment = 1
    
    company = random.choice(COMPANIES)
    doc.add_paragraph(f"Company: {company}")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d.%m.%Y')}")
    doc.add_paragraph()
    
    for i in range(random.randint(2, 4)):
        doc.add_heading(f"Section {i+1}: {random_sentence(2, 4)}", 1)
        for _ in range(random.randint(2, 5)):
            doc.add_paragraph(random_paragraph())
    
    doc.add_heading("Data Table", 2)
    table_data = random_table(random.randint(3, 6), 4)
    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    table.style = 'Table Grid'
    
    for i, row in enumerate(table_data):
        for j, cell in enumerate(row):
            table.rows[i].cells[j].text = str(cell)
    
    doc.save(path)
    print(f"Created: {path}")
    return path


def create_xlsx(path, title=None):
    """Create XLSX document."""
    if title is None:
        title = random.choice(TOPICS)
    
    wb = Workbook()
    ws1 = wb.active
    ws1.title = "Data"
    
    ws1['A1'] = title
    ws1['A1'].font = Font(size=16, bold=True)
    ws1.merge_cells('A1:E1')
    
    company = random.choice(COMPANIES)
    ws1['A2'] = f"Company: {company}"
    ws1['A3'] = f"Date: {datetime.now().strftime('%d.%m.%Y')}"
    
    table_data = random_table(random.randint(5, 10), 5)
    start_row = 5
    for i, row in enumerate(table_data):
        for j, val in enumerate(row):
            cell = ws1.cell(row=start_row+i, column=j+1, value=val)
            if i == 0:
                cell.font = Font(bold=True)
                cell.alignment = Alignment(horizontal='center')
    
    # Adjust column widths
    from openpyxl.utils import get_column_letter
    for col in range(1, len(table_data[0])+1):
        col_letter = get_column_letter(col)
        max_len = max(len(str(ws1.cell(row=start_row+r, column=col).value or "")) 
                     for r in range(len(table_data)))
        ws1.column_dimensions[col_letter].width = min(max_len + 2, 50)
    
    # Summary sheet
    ws2 = wb.create_sheet("Summary")
    ws2['A1'] = "Summary Report"
    ws2['A1'].font = Font(size=14, bold=True)
    
    summary = [
        ["Metric", "Value"],
        ["Total Items", random.randint(100, 1000)],
        ["Average Score", f"{random.randint(50, 100)}%"],
        ["Status", random.choice(["Active", "Pending", "Completed"])]
    ]
    
    for i, row in enumerate(summary):
        for j, val in enumerate(row):
            cell = ws2.cell(row=i+3, column=j+1, value=val)
            if i == 0:
                cell.font = Font(bold=True)
    
    wb.save(path)
    print(f"Created: {path}")
    return path


def create_pdf(path, title=None):
    """Create PDF document."""
    if title is None:
        title = random.choice(TOPICS)
    
    doc = SimpleDocTemplate(path, pagesize=letter,
                          rightMargin=72, leftMargin=72,
                          topMargin=72, bottomMargin=18)
    
    story = []
    styles = getSampleStyleSheet()
    
    title_style = styles['Heading1']
    title_style.alignment = 1
    story.append(Paragraph(title, title_style))
    story.append(Spacer(1, 12))
    
    normal_style = styles['Normal']
    company = random.choice(COMPANIES)
    story.append(Paragraph(f"<b>Company:</b> {company}", normal_style))
    story.append(Paragraph(f"<b>Date:</b> {datetime.now().strftime('%d.%m.%Y')}", normal_style))
    story.append(Spacer(1, 24))
    
    for i in range(random.randint(2, 4)):
        story.append(Paragraph(f"Section {i+1}: {random_sentence(2, 4)}", styles['Heading2']))
        story.append(Spacer(1, 12))
        for _ in range(random.randint(2, 4)):
            story.append(Paragraph(random_paragraph(), normal_style))
            story.append(Spacer(1, 12))
    
    doc.build(story)
    print(f"Created: {path}")
    return path


def create_zip(files, output_path):
    """Create ZIP archive."""
    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zf:
        for file_path in files:
            if os.path.exists(file_path):
                zf.write(file_path, os.path.basename(file_path))
    print(f"Created: {output_path}")
    return output_path


def create_7z(files, output_path):
    """Create 7z archive."""
    if not SEVENZIP_AVAILABLE:
        print("Warning: py7zr not available, skipping 7z")
        return None
    
    with py7zr.SevenZipFile(output_path, 'w') as archive:
        for file_path in files:
            if os.path.exists(file_path):
                archive.write(file_path, os.path.basename(file_path))
    print(f"Created: {output_path}")
    return output_path


def generate_all(output_dir="storage"):
    """Generate all test documents and archives."""
    docs_dir = Path(output_dir) / "documents"
    archives_dir = Path(output_dir) / "archives"
    
    docs_dir.mkdir(parents=True, exist_ok=True)
    archives_dir.mkdir(parents=True, exist_ok=True)
    
    created = {"docx": [], "xlsx": [], "pdf": [], "archives": []}
    
    print("=" * 60)
    print("Generating test documents...")
    print("=" * 60)
    
    # Generate documents
    for i in range(random.randint(3, 4)):
        title = random.choice(TOPICS)
        created["docx"].append(create_docx(str(docs_dir / f"document_{i+1}_{title.replace(' ', '_')}.docx"), title))
    
    for i in range(random.randint(3, 4)):
        title = random.choice(TOPICS)
        created["xlsx"].append(create_xlsx(str(docs_dir / f"spreadsheet_{i+1}_{title.replace(' ', '_')}.xlsx"), title))
    
    for i in range(random.randint(3, 4)):
        title = random.choice(TOPICS)
        created["pdf"].append(create_pdf(str(docs_dir / f"report_{i+1}_{title.replace(' ', '_')}.pdf"), title))
    
    # Create archives
    print("\n--- Creating archives ---")
    all_docs = created["docx"] + created["xlsx"] + created["pdf"]
    
    if len(all_docs) >= 3:
        # ZIP archive
        files_for_zip = random.sample(all_docs, min(3, len(all_docs)))
        zip_path = archives_dir / "documents_archive.zip"
        created["archives"].append(create_zip(files_for_zip, str(zip_path)))
        
        # 7z archive
        if SEVENZIP_AVAILABLE:
            files_for_7z = random.sample(all_docs, min(3, len(all_docs)))
            sevenzip_path = archives_dir / "documents_archive.7z"
            created["archives"].append(create_7z(files_for_7z, str(sevenzip_path)))
        
        # Nested archive
        if len(created["archives"]) > 0 and len(all_docs) >= 2:
            nested_files = random.sample(all_docs, 2)
            nested_zip = archives_dir / "nested_docs.zip"
            create_zip(nested_files, str(nested_zip))
            
            outer_zip = archives_dir / "nested_archive.zip"
            with zipfile.ZipFile(str(outer_zip), 'w', zipfile.ZIP_DEFLATED) as zf:
                zf.write(str(nested_zip), "nested_docs.zip")
                if all_docs:
                    zf.write(random.choice(all_docs), os.path.basename(all_docs[0]))
            
            created["archives"].append(str(outer_zip))
            print(f"Created nested archive: {outer_zip}")
            os.remove(str(nested_zip))
    
    print("\n" + "=" * 60)
    print("Generation complete!")
    print(f"DOCX: {len(created['docx'])}")
    print(f"XLSX: {len(created['xlsx'])}")
    print(f"PDF: {len(created['pdf'])}")
    print(f"Archives: {len(created['archives'])}")
    print(f"Total: {sum(len(v) for v in created.values())}")
    
    return created


if __name__ == "__main__":
    generate_all("storage")
