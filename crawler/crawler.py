#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Document Crawler - extracts text from documents and archives."""

import os
import csv
import tempfile
import shutil
import hashlib
from pathlib import Path
from datetime import datetime
from typing import Optional

from docx import Document
import openpyxl
import xlrd
import PyPDF2
import pdfplumber
import zipfile

try:
    import py7zr
    SEVENZIP_AVAILABLE = True
except ImportError:
    SEVENZIP_AVAILABLE = False
    print("Warning: py7zr not available. 7z archives will not be processed.")

try:
    import rarfile
    RARFILE_AVAILABLE = True
except ImportError:
    RARFILE_AVAILABLE = False
    print("Warning: rarfile not available. RAR archives will not be processed.")

SUPPORTED_EXTENSIONS = {
    'documents': ['.docx', '.doc', '.pdf'],
    'spreadsheets': ['.xlsx', '.xls'],
    'archives': ['.zip', '.7z', '.rar']
}

CSV_COLUMNS = [
    'id', 'file_path', 'file_name', 'file_type', 'file_size',
    'content', 'archive_path', 'created_date', 'content_hash'
]


def get_file_hash(file_path: str) -> str:
    """Calculate MD5 hash of a file."""
    hash_md5 = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_md5.update(chunk)
    return hash_md5.hexdigest()


def get_file_size(file_path: str) -> int:
    """Get file size in bytes."""
    return os.path.getsize(file_path)


def get_file_dates(file_path: str):
    """Get file creation and modification dates."""
    stat = os.stat(file_path)
    try:
        creation_time = stat.st_birthtime
    except AttributeError:
        creation_time = stat.st_ctime
    modification_time = stat.st_mtime
    return datetime.fromtimestamp(creation_time).isoformat(), \
           datetime.fromtimestamp(modification_time).isoformat()


def clean_text(text: str) -> str:
    """Clean extracted text."""
    if not text:
        return ""
    import re
    return re.sub(r'\s+', ' ', text).strip().replace('\x00', '')


def parse_docx(file_path: str) -> str:
    """Extract text from DOCX file."""
    try:
        doc = Document(file_path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return clean_text("\n".join(parts))
    except Exception as e:
        print(f"Error parsing DOCX {file_path}: {e}")
        return f"[Error parsing DOCX: {e}]"


def parse_xlsx(file_path: str) -> str:
    """Extract text from XLSX file."""
    try:
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            parts.append(f"[Sheet: {sheet_name}]")
            sheet = wb[sheet_name]
            for row in sheet.iter_rows(values_only=True):
                row_vals = [str(cell) for cell in row if cell is not None]
                if row_vals:
                    parts.append(" | ".join(row_vals))
        wb.close()
        return clean_text("\n".join(parts))
    except Exception as e:
        print(f"Error parsing XLSX {file_path}: {e}")
        return f"[Error parsing XLSX: {e}]"


def parse_xls(file_path: str) -> str:
    """Extract text from XLS file."""
    try:
        wb = xlrd.open_workbook(file_path)
        parts = []
        for sheet in wb.sheets():
            parts.append(f"[Sheet: {sheet.name}]")
            for row_idx in range(sheet.nrows):
                row_vals = [str(sheet.cell_value(row_idx, col_idx))
                           for col_idx in range(sheet.ncols)
                           if sheet.cell_value(row_idx, col_idx) is not None]
                if row_vals:
                    parts.append(" | ".join(row_vals))
        return clean_text("\n".join(parts))
    except Exception as e:
        print(f"Error parsing XLS {file_path}: {e}")
        return f"[Error parsing XLS: {e}]"


def parse_pdf_pypdf2(file_path: str) -> str:
    """Extract text from PDF using PyPDF2."""
    try:
        parts = []
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text.strip():
                    parts.append(f"[Page {page_num + 1}]")
                    parts.append(text)
        return clean_text("\n".join(parts))
    except Exception as e:
        print(f"Error parsing PDF with PyPDF2 {file_path}: {e}")
        return ""


def parse_pdf_pdfplumber(file_path: str) -> str:
    """Extract text from PDF using pdfplumber."""
    try:
        parts = []
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text()
                if text and text.strip():
                    parts.append(f"[Page {page_num}]")
                    parts.append(text)
        return clean_text("\n".join(parts))
    except Exception as e:
        print(f"Error parsing PDF with pdfplumber {file_path}: {e}")
        return ""


def parse_pdf(file_path: str) -> str:
    """Extract text from PDF, trying pdfplumber first then PyPDF2."""
    text = parse_pdf_pdfplumber(file_path)
    if len(text) < 50:
        backup = parse_pdf_pypdf2(file_path)
        if len(backup) > len(text):
            text = backup
    return text or "[No text could be extracted from PDF]"


def parse_document(file_path: str, file_ext: str) -> str:
    """Parse document and extract text."""
    parsers = {
        '.docx': parse_docx,
        '.doc': parse_docx,
        '.xlsx': parse_xlsx,
        '.xls': parse_xls,
        '.pdf': parse_pdf
    }
    parser = parsers.get(file_ext)
    return parser(file_path) if parser else f"[Unsupported file type: {file_ext}]"


def extract_zip(archive_path: str, extract_dir: str):
    """Extract ZIP archive."""
    extracted = []
    try:
        with zipfile.ZipFile(archive_path, 'r') as zf:
            zf.extractall(extract_dir)
            for name in zf.namelist():
                path = os.path.join(extract_dir, name)
                if os.path.isfile(path):
                    extracted.append(path)
    except Exception as e:
        print(f"Error extracting ZIP {archive_path}: {e}")
    return extracted


def extract_7z(archive_path: str, extract_dir: str):
    """Extract 7z archive."""
    if not SEVENZIP_AVAILABLE:
        return []
    extracted = []
    try:
        with py7zr.SevenZipFile(archive_path, 'r') as archive:
            archive.extractall(extract_dir)
            for name in archive.getnames():
                path = os.path.join(extract_dir, name)
                if os.path.isfile(path):
                    extracted.append(path)
    except Exception as e:
        print(f"Error extracting 7z {archive_path}: {e}")
    return extracted


def extract_rar(archive_path: str, extract_dir: str):
    """Extract RAR archive."""
    if not RARFILE_AVAILABLE:
        return []
    extracted = []
    try:
        with rarfile.RarFile(archive_path, 'r') as rf:
            rf.extractall(extract_dir)
            for name in rf.namelist():
                path = os.path.join(extract_dir, name)
                if os.path.isfile(path):
                    extracted.append(path)
    except Exception as e:
        print(f"Error extracting RAR {archive_path}: {e}")
    return extracted


def extract_archive(archive_path: str, extract_dir: str):
    """Extract archive based on extension."""
    ext = os.path.splitext(archive_path)[1].lower()
    extractors = {'.zip': extract_zip, '.7z': extract_7z, '.rar': extract_rar}
    extractor = extractors.get(ext)
    if extractor:
        return extractor(archive_path, extract_dir)
    print(f"Unsupported archive format: {ext}")
    return []


class DocumentCrawler:
    """Crawler for extracting text from documents and archives."""
    
    def __init__(self, storage_path: str, output_csv: str = "output/extracted_data.csv"):
        self.storage_path = Path(storage_path)
        self.output_csv = Path(output_csv)
        self.documents = []
        self.temp_dirs = []
    
    def _process_file(self, file_path: str, archive_path: str = "") -> Optional[dict]:
        """Process a single file."""
        file_ext = os.path.splitext(file_path)[1].lower()
        file_name = os.path.basename(file_path)
        
        supported = SUPPORTED_EXTENSIONS['documents'] + SUPPORTED_EXTENSIONS['spreadsheets']
        if file_ext not in supported:
            return None
        
        print(f"  Processing: {file_name}")
        
        file_size = get_file_size(file_path)
        created_date, _ = get_file_dates(file_path)
        content_hash = get_file_hash(file_path)
        content = parse_document(file_path, file_ext)
        
        file_type = 'document' if file_ext in SUPPORTED_EXTENSIONS['documents'] else 'spreadsheet'
        
        return {
            'id': len(self.documents) + 1,
            'file_path': file_path,
            'file_name': file_name,
            'file_type': file_type,
            'file_size': file_size,
            'content': content,
            'archive_path': archive_path,
            'created_date': created_date,
            'content_hash': content_hash
        }
    
    def _process_archive(self, archive_path: str, parent_path: str = "") -> list:
        """Process archive and its contents recursively."""
        results = []
        archive_name = os.path.basename(archive_path)
        full_path = f"{parent_path}/{archive_name}" if parent_path else archive_name
        print(f"\n  Extracting: {full_path}")
        
        temp_dir = tempfile.mkdtemp(prefix="crawler_")
        self.temp_dirs.append(temp_dir)
        
        for extracted in extract_archive(archive_path, temp_dir):
            ext = os.path.splitext(extracted)[1].lower()
            if ext in SUPPORTED_EXTENSIONS['archives']:
                results.extend(self._process_archive(extracted, full_path))
            else:
                doc = self._process_file(extracted, full_path)
                if doc:
                    results.append(doc)
        
        return results
    
    def crawl(self) -> list:
        """Scan storage and extract content from all documents."""
        print("=" * 60)
        print("Starting crawler...")
        print(f"Storage: {self.storage_path}")
        print(f"Output: {self.output_csv}")
        print()
        
        if not self.storage_path.exists():
            print(f"Error: Storage path does not exist")
            return []
        
        for root, _, files in os.walk(self.storage_path):
            for file_name in files:
                file_path = os.path.join(root, file_name)
                ext = os.path.splitext(file_path)[1].lower()
                
                print(f"\nFound: {file_name}")
                
                if ext in SUPPORTED_EXTENSIONS['archives']:
                    self.documents.extend(self._process_archive(file_path))
                else:
                    doc = self._process_file(file_path)
                    if doc:
                        self.documents.append(doc)
        
        self._cleanup()
        
        print("\n" + "=" * 60)
        print(f"Crawling complete! Processed {len(self.documents)} documents")
        
        return self.documents
    
    def _cleanup(self):
        """Remove temporary directories."""
        for temp_dir in self.temp_dirs:
            try:
                shutil.rmtree(temp_dir)
            except Exception as e:
                print(f"Warning: Could not remove {temp_dir}: {e}")
        self.temp_dirs.clear()
    
    def export(self) -> str:
        """Export data to CSV."""
        self.output_csv.parent.mkdir(parents=True, exist_ok=True)
        print(f"\nExporting to: {self.output_csv}")
        
        with open(self.output_csv, 'w', newline='', encoding='utf-8') as f:
            writer = csv.DictWriter(f, fieldnames=CSV_COLUMNS)
            writer.writeheader()
            writer.writerows(self.documents)
        
        print(f"Exported {len(self.documents)} documents")
        return str(self.output_csv)
    
    def run(self) -> str:
        """Run complete crawler pipeline."""
        self.crawl()
        return self.export()


def main():
    """Entry point."""
    import argparse
    
    parser = argparse.ArgumentParser(description='Document Crawler')
    parser.add_argument('--storage', '-s', default='storage', help='Storage directory')
    parser.add_argument('--output', '-o', default='output/extracted_data.csv', help='Output CSV')
    
    args = parser.parse_args()
    crawler = DocumentCrawler(args.storage, args.output)
    csv_path = crawler.run()
    print(f"\nDone! Output saved to: {csv_path}")
    return 0


if __name__ == "__main__":
    import sys
    sys.exit(main())
